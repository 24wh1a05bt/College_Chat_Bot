"""
Phase 1 - Ingest and index.

Loads data/college_website.docx, splits it by heading (so a chunk never
straddles two unrelated sections), then further splits long sections with
a RecursiveCharacterTextSplitter, embeds every chunk via OpenRouter, and
persists everything to a local ChromaDB collection.

Run directly to (re)build the index:
    python -m src.ingest --rebuild
"""
from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass, field

import chromadb
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from langchain_text_splitters import RecursiveCharacterTextSplitter

from src import config
from src.embeddings import OpenRouterEmbeddingFunction

# Word-per-page approximation used only to produce a human-friendly citation.
# A .docx has no fixed pagination in its XML (pagination is computed at
# render time), so we approximate "page N" from a running word count.
# This is documented in spec.md as a known limitation.
WORDS_PER_PAGE = 400


@dataclass
class Section:
    heading: str
    level: int
    paragraphs: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(p for p in self.paragraphs if p.strip())


def _table_to_text(table) -> str:
    """Convert a docx table into a text representation with rows and columns."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _find_tables_for_section(doc: DocxDocument, section_start_idx: int, section_end_idx: int) -> list[str]:
    """Find tables that appear between paragraph indices section_start_idx and section_end_idx.
    
    Tables in python-docx are stored separately from paragraphs, but we can
    determine which section a table belongs to by checking which paragraph
    it appears near. We iterate through the document's body XML elements in
    order to find tables that fall within the section's paragraph range.
    """
    tables_text = []
    # Get the body element
    body = doc.element.body
    # Track paragraph index as we walk XML elements
    para_idx = 0
    table_idx = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # It's a paragraph
            if section_start_idx <= para_idx < section_end_idx:
                pass  # within section
            para_idx += 1
        elif tag == 'tbl':
            # It's a table - check if it falls within our section range
            if section_start_idx <= para_idx <= section_end_idx + 1:
                if table_idx < len(doc.tables):
                    table_text = _table_to_text(doc.tables[table_idx])
                    if table_text.strip():
                        tables_text.append(table_text)
            table_idx += 1
    return tables_text


def parse_sections(doc: DocxDocument) -> list[Section]:
    """Walk the document and group paragraphs under their nearest heading.
    Also extracts table content that falls within each section."""
    sections: list[Section] = []
    current = Section(heading="Preamble", level=0)
    # Track paragraph indices for table-to-section mapping
    para_indices: list[tuple[int, int]] = []  # (start_idx, end_idx) per section

    current_start_idx = 0
    para_idx = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            para_idx += 1
            continue

        if style_name.startswith("heading") or style_name in ("title",):
            # Start a new section on every heading paragraph.
            if current.text:
                sections.append(current)
                para_indices.append((current_start_idx, para_idx))
            level = 1
            if style_name.startswith("heading"):
                digits = "".join(ch for ch in style_name if ch.isdigit())
                level = int(digits) if digits else 1
            current = Section(heading=text, level=level)
            current_start_idx = para_idx
        else:
            current.paragraphs.append(text)
        para_idx += 1

    if current.text:
        sections.append(current)
        para_indices.append((current_start_idx, para_idx))

    # Now extract tables for each section and append them as text
    for i, section in enumerate(sections):
        if i < len(para_indices):
            start_idx, end_idx = para_indices[i]
            table_texts = _find_tables_for_section(doc, start_idx, end_idx)
            for tt in table_texts:
                section.paragraphs.append(f"\n[Table Data]\n{tt}")

    # Drop an empty leading "Preamble" section if nothing landed in it.
    return [s for s in sections if s.text.strip()]


def _nearest_caption(paras: list, index: int, current_section_heading: str) -> str:
    """Many converted docs put the image in its own empty paragraph and the
    descriptive text in a sibling paragraph. Look one paragraph after, then
    one before, and use whichever has text (skipping headings so we don't
    accidentally grab the section title as a caption)."""
    for j in (index + 1, index - 1):
        if 0 <= j < len(paras):
            neighbor = paras[j]
            style_name = (neighbor.style.name or "").lower()
            neighbor_text = neighbor.text.strip()
            if neighbor_text and not (style_name.startswith("heading") or style_name == "title"):
                return neighbor_text
    return ""


def extract_images(doc: DocxDocument, images_dir) -> list[dict]:
    """Pull every embedded image out of the docx and tag it with the section
    it appears under and the best available caption text.

    Returns a list of {"filename", "section", "caption"} dicts. Images are
    written to disk at images_dir so the app can render them later - only
    the filename (not the bytes) is stored in Chroma metadata.
    """
    images_dir.mkdir(parents=True, exist_ok=True)
    records: list[dict] = []
    current_section = "Preamble"
    img_counter = 0
    paras = doc.paragraphs

    for i, para in enumerate(paras):
        style_name = (para.style.name or "").lower()
        text = para.text.strip()

        if style_name.startswith("heading") or style_name in ("title",):
            if text:
                current_section = text

        blips = para._element.findall(".//" + qn("a:blip"))
        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if not rId:
                continue
            try:
                image_part = doc.part.related_parts[rId]
            except KeyError:
                continue

            img_counter += 1
            content_type = getattr(image_part, "content_type", "image/png")
            ext = content_type.split("/")[-1].lower()
            ext = "jpg" if ext in ("jpeg", "jpg") else ext
            filename = f"image-{img_counter:03d}.{ext}"

            with open(images_dir / filename, "wb") as f:
                f.write(image_part.blob)

            records.append(
                {
                    "filename": filename,
                    "section": current_section,
                    "caption": text or _nearest_caption(paras, i, current_section),
                }
            )

    return records


def chunk_sections(sections: list[Section], image_records: list[dict] | None = None) -> list[dict]:
    """Split each section into overlapping chunks, carrying rich metadata."""
    image_records = image_records or []
    images_by_section: dict[str, list[dict]] = {}
    for rec in image_records:
        images_by_section.setdefault(rec["section"], []).append(
            {"filename": rec["filename"], "caption": rec.get("caption", "")}
        )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = []
    running_words = 0
    chunk_id = 0

    for section in sections:
        section_chunks = splitter.split_text(section.text)
        # Chroma metadata values must be str/int/float/bool, so a section's
        # images (there are usually 0-2) are serialized as a JSON string and
        # parsed back out in rag_chain.retrieve().
        section_images = images_by_section.get(section.heading, [])
        images_json = json.dumps(section_images) if section_images else ""

        for idx, chunk_text in enumerate(section_chunks):
            running_words += len(chunk_text.split())
            page = max(1, running_words // WORDS_PER_PAGE + 1)
            chunks.append(
                {
                    "id": f"chunk-{chunk_id:04d}",
                    "text": chunk_text,
                    "metadata": {
                        "source": config.DOCX_PATH.name,
                        "section": section.heading,
                        "section_chunk_index": idx,
                        "page": page,
                        "images": images_json,
                    },
                }
            )
            chunk_id += 1

    return chunks


def build_index(rebuild: bool = False) -> chromadb.Collection:
    if not config.DOCX_PATH.exists():
        sys.exit(
            f"Could not find {config.DOCX_PATH}. Place your college "
            f"document there (or set DOCX_PATH in .env)."
        )

    client = chromadb.PersistentClient(path=config.PERSIST_DIR)
    embed_fn = OpenRouterEmbeddingFunction()

    if rebuild:
        try:
            client.delete_collection(config.COLLECTION_NAME)
        except Exception:
            pass

    collection = client.get_or_create_collection(
        name=config.COLLECTION_NAME,
        embedding_function=embed_fn,
        metadata={"hnsw:space": "cosine"},
    )

    if collection.count() > 0 and not rebuild:
        print(f"Index already has {collection.count()} chunks. Use --rebuild to re-index.")
        return collection

    doc = DocxDocument(str(config.DOCX_PATH))
    sections = parse_sections(doc)
    image_records = extract_images(doc, config.IMAGES_DIR)
    chunks = chunk_sections(sections, image_records)

    if not chunks:
        sys.exit("No chunkable content found in the document.")

    collection.add(
        ids=[c["id"] for c in chunks],
        documents=[c["text"] for c in chunks],
        metadatas=[c["metadata"] for c in chunks],
    )

    print(f"Indexed {len(chunks)} chunks from {len(sections)} sections into "
          f"'{config.COLLECTION_NAME}' at {config.PERSIST_DIR}")
    return collection


def get_collection() -> chromadb.Collection:
    """Open the existing collection without re-embedding (used by the app)."""
    client = chromadb.PersistentClient(path=config.PERSIST_DIR)
    return client.get_or_create_collection(
        name=config.COLLECTION_NAME,
        embedding_function=OpenRouterEmbeddingFunction(),
        metadata={"hnsw:space": "cosine"},
    )


def list_sections() -> list[str]:
    """Distinct section names currently in the index, for the sidebar filter."""
    collection = get_collection()
    if collection.count() == 0:
        return config.KNOWN_SECTIONS
    got = collection.get(include=["metadatas"])
    sections = sorted({m["section"] for m in got["metadatas"]})
    return sections


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--rebuild", action="store_true", help="Force re-indexing")
    args = parser.parse_args()
    build_index(rebuild=args.rebuild)