# """
# College website scraper -> ONE Word document, with each page as its own
# section (heading + page break), including tables and images in reading
# order.

# Usage:
#     pip install requests beautifulsoup4 python-docx pillow --break-system-packages
#     python scrape_single_doc.py
# """

# import io
# import time
# from collections import deque
# from urllib.parse import urljoin, urlparse

# import requests
# from bs4 import BeautifulSoup
# from docx import Document
# from docx.shared import Inches
# from PIL import Image

# START_URL = "https://bvrithyderabad.edu.in/"
# OUTPUT_FILE = "college_website.docx"
# MAX_PAGES = 200
# REQUEST_DELAY = 0.5
# TIMEOUT = 20
# MAX_IMAGE_BYTES = 8 * 1024 * 1024   # skip anything larger than 8MB
# MIN_IMAGE_DIM = 60                  # skip icons/tracking pixels smaller than this

# HEADERS = {
#     "User-Agent": "Mozilla/5.0 (compatible; SiteArchiveBot/1.0; +https://example.com/bot)"
# }


# def clean_url(url: str) -> str:
#     parsed = urlparse(url)
#     return f"{parsed.scheme}://{parsed.netloc}{parsed.path}"


# def is_probably_document_link(url: str) -> bool:
#     skip_ext = (
#         ".pdf", ".zip", ".mp4", ".mp3", ".doc", ".docx", ".xls", ".xlsx",
#         ".ppt", ".pptx", ".css", ".js", ".ico",
#     )
#     return urlparse(url).path.lower().endswith(skip_ext)


# def add_table_to_doc(document: Document, table_tag) -> None:
#     rows = table_tag.find_all("tr")
#     if not rows:
#         return
#     row_cells = [r.find_all(["td", "th"]) for r in rows]
#     n_cols = max((len(cells) for cells in row_cells), default=0)
#     if n_cols == 0:
#         return

#     word_table = document.add_table(rows=0, cols=n_cols)
#     word_table.style = "Light Grid Accent 1"

#     for cells in row_cells:
#         widget_cells = word_table.add_row().cells
#         for i in range(n_cols):
#             text = cells[i].get_text(separator=" ", strip=True) if i < len(cells) else ""
#             widget_cells[i].text = text
#             if cells and i < len(cells) and cells[i].name == "th":
#                 for p in widget_cells[i].paragraphs:
#                     for run in p.runs:
#                         run.bold = True

#     document.add_paragraph()


# def add_image_to_doc(document: Document, img_tag, page_url: str, session: requests.Session) -> None:
#     """Download an <img> and embed it. Skips tiny/broken/oversized images."""
#     src = img_tag.get("src") or img_tag.get("data-src")
#     if not src:
#         return
#     if src.startswith("data:"):
#         return  # skip inline base64 images to keep things simple

#     absolute = urljoin(page_url, src)

#     try:
#         resp = session.get(absolute, timeout=TIMEOUT, stream=True)
#         if resp.status_code != 200:
#             return

#         content_length = resp.headers.get("Content-Length")
#         if content_length and int(content_length) > MAX_IMAGE_BYTES:
#             return

#         raw = resp.content
#         if len(raw) > MAX_IMAGE_BYTES or len(raw) == 0:
#             return

#         image_buffer = io.BytesIO(raw)
#         with Image.open(image_buffer) as im:
#             im.verify()
#         image_buffer.seek(0)
#         with Image.open(image_buffer) as im:
#             width, height = im.size
#             if width < MIN_IMAGE_DIM or height < MIN_IMAGE_DIM:
#                 return
#             fmt = im.format

#         image_buffer.seek(0)

#         if fmt == "WEBP" or fmt is None:
#             image_buffer.seek(0)
#             with Image.open(image_buffer) as im:
#                 converted = io.BytesIO()
#                 im.convert("RGB").save(converted, format="PNG")
#                 converted.seek(0)
#                 image_buffer = converted

#         document.add_picture(image_buffer, width=Inches(5.5))

#         alt = img_tag.get("alt", "").strip()
#         if alt:
#             caption = document.add_paragraph(alt)
#             if "Caption" in document.styles:
#                 caption.style = document.styles["Caption"]

#     except Exception:
#         return  # broken/unreachable/unsupported image — skip silently


# def render_page_body(document: Document, soup: BeautifulSoup, page_url: str, session: requests.Session) -> None:
#     body = soup.body or soup

#     for junk in body.find_all(["script", "style", "noscript", "header", "footer", "nav"]):
#         junk.decompose()

#     content_tags = {"p", "li", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "img", "table"}
#     seen_any_content = False

#     for element in body.find_all(list(content_tags), recursive=True):
#         if element.find_parent("table") is not None:
#             continue

#         if element.name == "table":
#             add_table_to_doc(document, element)
#             seen_any_content = True
#             continue

#         if element.name == "img":
#             add_image_to_doc(document, element, page_url, session)
#             seen_any_content = True
#             continue

#         text = element.get_text(separator=" ", strip=True)
#         if not text:
#             continue

#         if element.name.startswith("h") and element.name[1:].isdigit():
#             # offset by 1 so the page title (level 1) stays the top heading
#             # within this section; cap at 9 (Word's max heading level)
#             level = min(int(element.name[1]) + 1, 9)
#             document.add_heading(text, level=level)
#         else:
#             document.add_paragraph(text)

#         seen_any_content = True

#     if not seen_any_content:
#         document.add_paragraph("[No extractable text, tables, or images found on this page]")


# def main():
#     domain = urlparse(START_URL).netloc
#     visited = set()
#     queue = deque([clean_url(START_URL)])

#     document = Document()
#     document.add_heading("College Website Content", level=0)

#     pages_scraped = 0
#     session = requests.Session()
#     session.headers.update(HEADERS)

#     while queue and pages_scraped < MAX_PAGES:
#         url = queue.popleft()
#         if url in visited:
#             continue
#         visited.add(url)

#         try:
#             response = session.get(url, timeout=TIMEOUT)
#             if response.status_code != 200:
#                 continue
#             if "text/html" not in response.headers.get("Content-Type", ""):
#                 continue

#             html = response.text
#             soup = BeautifulSoup(html, "html.parser")

#             # --- start a new section for this page ---
#             if pages_scraped > 0:
#                 document.add_page_break()

#             title = soup.title.get_text(strip=True) if soup.title else url
#             document.add_heading(title, level=1)
#             url_para = document.add_paragraph(url)
#             url_para.italic = True

#             render_page_body(document, soup, url, session)

#             pages_scraped += 1
#             print(f"[{pages_scraped}] added section for: {url}")

#             for link in soup.find_all("a", href=True):
#                 absolute = urljoin(url, link["href"])
#                 candidate = clean_url(absolute)
#                 parsed = urlparse(candidate)

#                 if parsed.netloc != domain:
#                     continue
#                 if is_probably_document_link(candidate):
#                     continue
#                 if candidate not in visited:
#                     queue.append(candidate)

#             time.sleep(REQUEST_DELAY)

#         except Exception as e:
#             print(f"Error on {url}: {e}")

#     document.save(OUTPUT_FILE)
#     print(f"Finished. Scraped {pages_scraped} pages into {OUTPUT_FILE}")


# if __name__ == "__main__":
#     main()

#!/usr/bin/env python3
# """
# Scrape images from the BVRIT Hyderabad official website
# (https://bvrithyderabad.edu.in/) for use in the college knowledge-base document.

# Setup:
#     pip install requests beautifulsoup4

# Run:
#     python scrape_bvrith_images.py

# Output:
#     Downloaded images land in ./bvrith_images/, named after the page they
#     came from + the original filename (e.g. principal__PrincipalCSE.jpg).
#     A manifest.csv is also written with page, image URL, and local filename,
#     so you can see at a glance what you got and where it came from.
# """

# import os
# import csv
# import re
# import time
# from urllib.parse import urljoin, urlparse

# import requests
# from bs4 import BeautifulSoup

# BASE_URL = "https://bvrithyderabad.edu.in/"

# # Pages likely to contain the photos you want (leadership, principal, campus,
# # facilities). Add/remove URLs here to steer what gets scraped.
# PAGES_TO_SCRAPE = [
#     "",                     # homepage
#     "about-us/",
#     "chairman-message/",
#     "vice-chairman-message/",
#     "secretary-message/",
#     "principal/",
#     "principal-message/",
#     "management/",
#     "campus/",
#     "infrastructure/",
#     "library/",
#     "hostel/",
#     "sports/",
#     "facilities/",
#     "gallery/",
#     "photo-gallery/",
# ]

# # Skip obvious junk: tiny icons, logos, social-media badges, spacer gifs, etc.
# SKIP_PATTERNS = re.compile(
#     r"(logo|icon|favicon|sprite|social|arrow|bullet|spacer|placeholder)",
#     re.IGNORECASE,
# )

# MIN_BYTES = 15_000  # skip anything smaller than ~15KB (usually icons/logos)

# OUTPUT_DIR = "bvrith_images"

# HEADERS = {
#     # A normal browser User-Agent avoids some basic bot-blocking (403s).
#     "User-Agent": (
#         "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
#         "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
#     )
# }


# def get_image_urls(page_url, session):
#     """Return a set of absolute image URLs found on a page."""
#     try:
#         resp = session.get(page_url, headers=HEADERS, timeout=15)
#         resp.raise_for_status()
#     except requests.RequestException as e:
#         print(f"  [skip] could not load {page_url}: {e}")
#         return set()

#     soup = BeautifulSoup(resp.text, "html.parser")
#     urls = set()

#     for img in soup.find_all("img"):
#         # Prefer the highest-res candidate: srcset > data-src (lazy load) > src
#         candidates = []
#         if img.get("srcset"):
#             # srcset format: "url1 480w, url2 800w, ..." — take the last (largest)
#             parts = [p.strip().split(" ")[0] for p in img["srcset"].split(",")]
#             candidates.extend(parts)
#         for attr in ("data-src", "data-lazy-src", "src"):
#             if img.get(attr):
#                 candidates.append(img[attr])

#         for c in candidates:
#             full_url = urljoin(page_url, c)
#             if SKIP_PATTERNS.search(full_url):
#                 continue
#             if full_url.lower().endswith((".jpg", ".jpeg", ".png", ".webp")):
#                 urls.add(full_url)

#     return urls


# def download_image(url, page_slug, session, seen_filenames):
#     filename = os.path.basename(urlparse(url).path)
#     if not filename:
#         return None

#     local_name = f"{page_slug}__{filename}" if page_slug else filename
#     if local_name in seen_filenames:
#         return None  # already downloaded (e.g. same image on multiple pages)
#     seen_filenames.add(local_name)

#     local_path = os.path.join(OUTPUT_DIR, local_name)

#     try:
#         resp = session.get(url, headers=HEADERS, timeout=15)
#         resp.raise_for_status()
#     except requests.RequestException as e:
#         print(f"    [fail] {url}: {e}")
#         return None

#     if len(resp.content) < MIN_BYTES:
#         print(f"    [skip] too small, likely an icon: {url}")
#         return None

#     with open(local_path, "wb") as f:
#         f.write(resp.content)

#     print(f"    [ok] {url} -> {local_name} ({len(resp.content)//1024} KB)")
#     return local_path


# def main():
#     os.makedirs(OUTPUT_DIR, exist_ok=True)
#     session = requests.Session()
#     seen_filenames = set()
#     manifest_rows = []

#     for page_path in PAGES_TO_SCRAPE:
#         page_url = urljoin(BASE_URL, page_path)
#         slug = page_path.strip("/").replace("/", "-") or "home"
#         print(f"Scanning {page_url} ...")

#         image_urls = get_image_urls(page_url, session)
#         if not image_urls:
#             print("  (no images found or page unreachable)")
#             continue

#         for img_url in sorted(image_urls):
#             local_path = download_image(img_url, slug, session, seen_filenames)
#             if local_path:
#                 manifest_rows.append(
#                     {"page": page_url, "image_url": img_url, "local_file": local_path}
#                 )
#             time.sleep(0.2)  # be polite to the server

#     manifest_path = os.path.join(OUTPUT_DIR, "manifest.csv")
#     with open(manifest_path, "w", newline="", encoding="utf-8") as f:
#         writer = csv.DictWriter(f, fieldnames=["page", "image_url", "local_file"])
#         writer.writeheader()
#         writer.writerows(manifest_rows)

#     print(f"\nDone. {len(manifest_rows)} images saved to ./{OUTPUT_DIR}/")
#     print(f"See {manifest_path} for a full list of what was downloaded and where it came from.")


# if __name__ == "__main__":
#     main()


#!/usr/bin/env python3
"""
Add every image from a local folder (e.g. the bvrith_images/ folder produced
by scrape_bvrith_images.py) into an existing Word document, WITHOUT touching
any of the document's existing text/content.

It works by opening the .docx as-is and appending a new section at the end
titled "Additional Images", with each image inserted one after another
(captioned with its filename, or with page info if a manifest.csv is found
alongside the images). Everything already in the document is left exactly
as it was.

Setup:
    pip install python-docx

Run:
    python merge_images_into_doc.py college_website.docx bvrith_images/

    (first arg = your existing .docx, second arg = folder of images)

Output:
    college_website_with_images.docx  (a new file — your original is untouched)
"""

import csv
import os
import sys

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp")
MAX_WIDTH_INCHES = 5.5  # matches the image width already used in the document


def load_manifest(folder):
    """If manifest.csv (written by the scraper) is present, map filename -> page URL."""
    manifest_path = os.path.join(folder, "manifest.csv")
    mapping = {}
    if os.path.exists(manifest_path):
        with open(manifest_path, newline="", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                local_file = os.path.basename(row.get("local_file", ""))
                if local_file:
                    mapping[local_file] = row.get("page", "")
    return mapping


def caption_for(filename, page_map):
    """Build a human-readable caption from the filename / manifest page info."""
    name = os.path.splitext(filename)[0]
    # strip the "<pageslug>__" prefix added by the scraper, if present
    if "__" in name:
        page_slug, rest = name.split("__", 1)
    else:
        page_slug, rest = "", name

    pretty = rest.replace("-", " ").replace("_", " ").strip()
    caption = pretty if pretty else filename

    page_url = page_map.get(filename, "")
    if page_url:
        caption += f" (source: {page_url})"
    elif page_slug:
        caption += f" (from: {page_slug} page)"

    return caption


def main():
    if len(sys.argv) != 3:
        print("Usage: python merge_images_into_doc.py <existing.docx> <images_folder>")
        sys.exit(1)

    doc_path, images_folder = sys.argv[1], sys.argv[2]

    if not os.path.exists(doc_path):
        print(f"Document not found: {doc_path}")
        sys.exit(1)
    if not os.path.isdir(images_folder):
        print(f"Images folder not found: {images_folder}")
        sys.exit(1)

    doc = Document(doc_path)

    image_files = sorted(
        f for f in os.listdir(images_folder) if f.lower().endswith(IMAGE_EXTENSIONS)
    )
    if not image_files:
        print("No images found in that folder.")
        sys.exit(1)

    page_map = load_manifest(images_folder)

    # --- Append a new section; nothing above this point is touched ---
    doc.add_page_break()
    heading = doc.add_heading("Additional Images", level=1)

    added, skipped = 0, 0
    for filename in image_files:
        img_path = os.path.join(images_folder, filename)
        try:
            doc.add_picture(img_path, width=Inches(MAX_WIDTH_INCHES))
        except Exception as e:
            print(f"  [skip] {filename}: {e}")
            skipped += 1
            continue

        # Center the image
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add an italic caption under it
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption_for(filename, page_map))
        run.italic = True

        added += 1
        print(f"  [ok] added {filename}")

    out_path = os.path.splitext(doc_path)[0] + "_with_images.docx"
    doc.save(out_path)

    print(f"\nDone. {added} images added, {skipped} skipped.")
    print(f"Saved as: {out_path}")
    print("Your original document was not modified.")


if __name__ == "__main__":
    main()