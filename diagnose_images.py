"""
One-off diagnostic - run this from your project root:

    python diagnose_images.py

It tells us HOW images are embedded in data/college_website.docx (modern
DrawingML vs. legacy VML vs. not embedded at all), which determines the fix.
Paste the full output back into the chat.
"""
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from src import config

doc = DocxDocument(str(config.DOCX_PATH))

print(f"Document: {config.DOCX_PATH}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"python-docx inline_shapes count: {len(doc.inline_shapes)}")
print()

blip_count = 0
vml_count = 0
hyperlink_img_count = 0

print("--- Paragraphs containing image-like XML, with a snippet of nearby text ---")
current_heading = "Preamble"
for para in doc.paragraphs:
    style_name = (para.style.name or "").lower()
    text = para.text.strip()
    if style_name.startswith("heading") or style_name in ("title",):
        if text:
            current_heading = text

    blips = para._element.findall(".//" + qn("a:blip"))
    # legacy VML fallback images (common in "save webpage as .docx" exports)
    vml = para._element.findall(".//{urn:schemas-microsoft-com:vml}imagedata")
    # some converters store the image only as an OLE object / hyperlink
    has_drawing = para._element.findall(".//" + qn("w:drawing"))
    has_pict = para._element.findall(".//" + qn("w:pict"))

    if blips or vml or has_pict:
        blip_count += len(blips)
        vml_count += len(vml)
        print(
            f"[{current_heading}] blips={len(blips)} vml_imagedata={len(vml)} "
            f"w:pict={len(has_pict)} w:drawing={len(has_drawing)} | text={text[:80]!r}"
        )

print()
print("--- Sample raw text under the 'Campus Banners' style section (first 15 paragraphs) ---")
capturing = False
shown = 0
for para in doc.paragraphs:
    style_name = (para.style.name or "").lower()
    text = para.text.strip()
    if style_name.startswith("heading") or style_name in ("title",):
        capturing = "campus" in text.lower() and "banner" in text.lower() or "scrolling" in text.lower()
        if capturing:
            print(f"HEADING: {text}")
        continue
    if capturing and text and shown < 15:
        print(f"  para: {text[:120]!r}")
        shown += 1

print()
print("--- Relationship parts on the main document (rId -> content type) ---")
for rel_id, rel in doc.part.rels.items():
    if "image" in rel.reltype.lower():
        print(f"  {rel_id}: {rel.reltype} -> target_ref={rel.target_ref}")

print()
print(f"TOTALS: a:blip={blip_count}  v:imagedata={vml_count}")